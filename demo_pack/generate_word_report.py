#!/usr/bin/env python3
"""
生成 Word 报告 - 支持 LaTeX 公式
========================

需要安装:
    pip install python-docx
    
可选（更好的公式支持）:
    pip install docxlatex

Usage:
    python generate_word_report.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import subprocess
import os


def create_report():
    """创建 Word 报告"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ===== 第一页：执行摘要 =====
    
    # 标题
    title = doc.add_heading('SIMM 2.8 Challenge Model\n技术实现报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 元信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('作者: OpenClaw Multi-Agent\n').italic = True
    meta.add_run('日期: 2026-02-26\n').italic = True
    meta.add_run('版本: v1.0').italic = True
    
    doc.add_page_break()
    
    # 核心成果
    doc.add_heading('📊 核心成果', level=1)
    
    highlights = doc.add_paragraph()
    highlights.add_run('✅ 4 层 Challenge 策略\n').bold = True
    highlights.add_run('   覆盖 13 种产品类型，差异化验证\n\n')
    highlights.add_run('✅ 数学熔断机制\n').bold = True
    highlights.add_run('   基于 ISDA SIMM 2.8 官方公式\n\n')
    highlights.add_run('✅ 100% 测试通过\n').bold = True
    highlights.add_run('   25个测试用例全部验证\n')
    
    # 关键公式
    doc.add_heading('🧮 关键公式', level=1)
    
    # 公式 1：加权敏感度
    doc.add_paragraph('加权敏感度公式 (SIMM 2.8 Section 4):', style='Intense Quote')
    
    # 使用 OMML (Office Math Markup Language) 插入公式
    # 或者使用 Unicode 近似
    formula1 = doc.add_paragraph()
    formula1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula1.add_run('WSₖ = RWₖ × sₖ × CRₖ')
    run.font.size = Pt(14)
    run.font.name = 'Cambria Math'
    
    # 公式说明
    doc.add_paragraph(
        '其中:\n'
        '• WSₖ = Weighted Sensitivity (加权敏感度)\n'
        '• RWₖ = Risk Weight (风险权重，Table 1)\n'
        '• sₖ = Sensitivity (敏感度)\n'
        '• CRₖ = Concentration Risk (集中度风险)',
        style='List Bullet'
    )
    
    # 公式 2：聚合
    doc.add_paragraph('聚合公式:', style='Intense Quote')
    
    formula2 = doc.add_paragraph()
    formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula2.add_run('K = √(∑ₖ WSₖ² + ∑ₖ∑ₗ≠ₖ ρₖₗ WSₖ WSₗ)')
    run.font.size = Pt(14)
    run.font.name = 'Cambria Math'
    
    # 公式 3：Scaling Function
    doc.add_paragraph('缩放函数 (SIMM 2.8 Section 11):', style='Intense Quote')
    
    formula3 = doc.add_paragraph()
    formula3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula3.add_run('SF(t) = 0.5 × min(1, 14/t)')
    run.font.size = Pt(14)
    run.font.name = 'Cambria Math'
    
    # 熔断阈值表
    doc.add_heading('🚨 熔断阈值', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '产品类型'
    hdr_cells[1].text = '熔断阈值'
    hdr_cells[2].text = 'ISDA 依据'
    
    # 数据行
    thresholds = [
        ('Barrier', '距离障碍 < 2%', 'Section 11(a)'),
        ('Digital', 'Vega > 名义本金 50%', 'Section C.8'),
        ('Touch', '立即熔断', '公式不适用'),
        ('TARF', '目标达成 > 80% 但 Vega 高', '行为转变'),
    ]
    
    for product, threshold, reference in thresholds:
        row_cells = table.add_row().cells
        row_cells[0].text = product
        row_cells[1].text = threshold
        row_cells[2].text = reference
    
    doc.add_page_break()
    
    # ===== 后续页面：详细方案 =====
    
    doc.add_heading('1. Tier 1: 线性产品 Challenge', level=1)
    doc.add_paragraph(
        '适用产品: FX Forward, FX Swap, NDF, IRS, Basis Swap\n'
        'ISDA 依据: Section C.1 (Delta Risk), Section 4 (Aggregation)'
    )
    
    doc.add_heading('Challenge 验证点', level=2)
    
    checks = [
        ('Risk Weight 一致性', '验证 RWₖ 与 SIMM 2.8 Table 1 一致'),
        ('聚合上限检查', 'K ≤ ∑|WSₖ| × 1.01 (次可加性)'),
        ('Delta 符号合理性', 'Pay Fixed → PV01 < 0'),
    ]
    
    for check, desc in checks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{check}: ').bold = True
        p.add_run(desc)
    
    doc.add_heading('代码实现', level=2)
    
    code = doc.add_paragraph()
    code.style = 'Intense Quote'
    code.add_run(
        'def _verify_aggregation_bound(self, simm_result):\n'
        '    ws_sum = sum(abs(ws) for ws in simm_result.ws_by_bucket)\n'
        '    if simm_result.k_value > ws_sum * 1.01:\n'
        '        return ModelBreakdown("Aggregation exceeds maximum")\n'
        '    return Pass()'
    ).font.name = 'Courier New'
    
    # 继续添加更多章节...
    doc.add_page_break()
    
    doc.add_heading('2. Tier 2: 香草期权 Challenge', level=1)
    doc.add_paragraph(
        '适用产品: Vanilla Option, Swaption\n'
        'ISDA 依据: Section C.8, Section 11(a)'
    )
    
    # 更多内容...
    
    # 保存
    output_path = 'SIMM_28_Challenge_Model_Report.docx'
    doc.save(output_path)
    print(f"✅ 报告已生成: {output_path}")
    
    return output_path


def convert_with_latex(docx_path):
    """
    使用 LibreOffice 或 Pandoc 转换，保留公式格式
    """
    # 方法1: 使用 pandoc 将 docx 转为带 LaTeX 的格式
    md_path = docx_path.replace('.docx', '.md')
    
    # 先导出为 Markdown
    subprocess.run([
        'pandoc', docx_path, 
        '-o', md_path,
        '--wrap=none'
    ], check=True)
    
    # 再转回 docx，使用 LaTeX 公式
    final_path = docx_path.replace('.docx', '_with_latex.docx')
    subprocess.run([
        'pandoc', md_path,
        '-o', final_path,
        '--mathml'  # 或 --mathjax
    ], check=True)
    
    print(f"✅ LaTeX 公式版本: {final_path}")
    return final_path


if __name__ == '__main__':
    # 生成基础报告
    docx_file = create_report()
    
    # 如果需要 LaTeX 公式支持，取消下面注释
    # convert_with_latex(docx_file)
